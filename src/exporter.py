import pandas as pd
from docx import Document
import io
from openpyxl.styles import PatternFill, Font

def parsear_tabla_markdown(texto_md):
    """Convierte la tabla markdown de la IA en una lista de diccionarios."""
    filas = []
    lineas = texto_md.strip().split('\n')
    for linea in lineas:
        if '|' in linea and 'Requisito Original' not in linea and '---' not in linea:
            columnas = [col.strip() for col in linea.split('|') if col.strip()]
            if len(columnas) >= 2:
                filas.append({"Requisito Original": columnas[0], "Dictamen Técnico": columnas[1]})
    return filas

def exportar_excel(texto_md, nombre_partida, equipo_nombre):
    datos = parsear_tabla_markdown(texto_md)
    df = pd.DataFrame(datos)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Escribimos los datos a partir de la fila 4 (dejamos espacio para los títulos)
        df.to_excel(writer, sheet_name='Dictamen', startrow=3, index=False)
        worksheet = writer.sheets['Dictamen']
        
        # Ajustamos el ancho de las columnas
        worksheet.column_dimensions['A'].width = 70
        worksheet.column_dimensions['B'].width = 90
        
        # --- ENCABEZADOS PERSONALIZADOS ---
        worksheet['A1'] = f"Dictamen Técnico: {nombre_partida}"
        worksheet['A1'].font = Font(bold=True, size=14)
        worksheet['A2'] = f"Equipo Evaluado: {equipo_nombre}"
        worksheet['A2'].font = Font(bold=True, size=12)
        
        # --- COLORES DE CELDAS (Código ARGB) ---
        fill_verde = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")   # Verde pastel
        fill_rojo = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")    # Rojo pastel
        fill_amarillo = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid") # Amarillo pastel
        fill_azul = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")    # Azul pastel
        fill_blanco = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")  # Gris claro
        
        # --- LÓGICA DE PINTADO ---
        # Recorremos la columna B (Dictamen) empezando desde la fila 5 (donde inician los datos)
        for row in range(5, len(datos) + 5):
            celda_dictamen = worksheet[f'B{row}']
            texto = str(celda_dictamen.value)
            
            # Pintamos dependiendo del emoji detectado
            if '🟢' in texto:
                celda_dictamen.fill = fill_verde
            elif '🔴' in texto or '🟥' in texto:
                celda_dictamen.fill = fill_rojo
            elif '🟡' in texto:
                celda_dictamen.fill = fill_amarillo
            elif '🔵' in texto:
                celda_dictamen.fill = fill_azul
            elif '⚪' in texto:
                celda_dictamen.fill = fill_blanco

    return output.getvalue()

def exportar_word(texto_md, nombre_partida, equipo_nombre):
    datos = parsear_tabla_markdown(texto_md)
    doc = Document()
    
    # --- ENCABEZADOS EN WORD ---
    doc.add_heading(f'Dictamen Técnico: {nombre_partida}', 0)
    doc.add_paragraph(f'Equipo Evaluado: {equipo_nombre}')
    
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = 'Requisito Original'
    hdr_cells[1].text = 'Dictamen Técnico'
    
    for item in datos:
        row_cells = tabla.add_row().cells
        row_cells[0].text = item['Requisito Original']
        row_cells[1].text = item['Dictamen Técnico']
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()